# -*- coding: utf-8 -*-
"""Convert docs/RESEARCH_OUTCOMES.md to a formatted Word document with embedded figures."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "docs" / "RESEARCH_OUTCOMES.md"
OUT = ROOT / "docs" / "Aadhaar_Intel_Engine_Research_Outcomes.docx"
DOCS = ROOT / "docs"


def resolve_img(src: str) -> Path | None:
    src = src.strip()
    candidates = [
        DOCS / src,
        ROOT / src,
        DOCS / Path(src).name,
        ROOT / "images" / Path(src).name,
        DOCS / "live_captures" / Path(src).name,
    ]
    for c in candidates:
        if c.exists() and c.is_file():
            return c
    if src.startswith("../"):
        p = (DOCS / src).resolve()
        if p.exists():
            return p
    return None


def set_run_font(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_header_row(row, fill="1E3A5F"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        set_run_font(run, size=9, bold=True, color=(255, 255, 255))
    shade_header_row(hdr, "1E3A5F")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val is not None else "")
            set_run_font(run, size=9)
            if ri % 2 == 1:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
                tcPr.append(shd)
    doc.add_paragraph()
    return t


def parse_table(lines, start_idx):
    header_line = lines[start_idx].strip()
    if start_idx + 1 >= len(lines):
        return None, None, start_idx + 1
    sep = lines[start_idx + 1].strip()
    if not re.match(r"^\|[\s\-:|]+\|$", sep):
        return None, None, start_idx + 1

    def split_row(line):
        line = line.strip().strip("|")
        return [c.strip() for c in line.split("|")]

    headers = split_row(header_line)
    rows = []
    i = start_idx + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(split_row(lines[i]))
        i += 1
    rows = [r[: len(headers)] + [""] * max(0, len(headers) - len(r)) for r in rows]
    return headers, rows, i


def add_image(doc, path: Path, caption: str | None = None, max_width=6.2):
    try:
        from PIL import Image

        with Image.open(path) as im:
            w, h = im.size
        aspect = h / w if w else 1
        width = max_width
        height = max_width * aspect
        if height > 7.5:
            height = 7.5
            width = height / aspect
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            set_run_font(r, size=9, italic=True, color=(71, 85, 105))
        return True
    except Exception:
        p = doc.add_paragraph()
        r = p.add_run(f"[Image not available: {path.name}]")
        set_run_font(r, size=9, italic=True, color=(148, 163, 184))
        return False


def add_inline_markdown(p, content: str, size=11):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", content)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_run_font(r, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            set_run_font(r, size=size - 1 if size > 9 else size, font="Consolas", color=(30, 64, 175))
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                r = p.add_run(m.group(1))
                set_run_font(r, size=size, color=(37, 99, 235))
                r.underline = True
            else:
                r = p.add_run(part)
                set_run_font(r, size=size)
        else:
            r = p.add_run(part)
            set_run_font(r, size=size)


def main():
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        st = styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st.font.color.rgb = RGBColor(30, 58, 95)
        st.font.bold = True
        st.font.size = Pt(size)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Aadhaar Intel Engine — Research Outcomes")
    set_run_font(hr, size=9, color=(100, 116, 139))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Research outcomes report · Page ")
    set_run_font(fr, size=9, color=(100, 116, 139))
    fld = parse_xml(r'<w:fldSimple %s w:instr=" PAGE "/>' % nsdecls("w"))
    fp._p.append(fld)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Aadhaar Intel Engine")
    set_run_font(tr, size=28, bold=True, color=(30, 58, 95))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Research Outcomes Report")
    set_run_font(sr, size=18, bold=True, color=(37, 99, 235))

    for m in [
        "UIDAI operational aggregates · Patterns, anomalies, and predictive indicators",
        "Code-aligned living document with figure and table provisions",
        "Repository: github.com/NetRunnerXD/Aadhaar-Intel-Engine-UIDAI",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(m)
        set_run_font(r, size=11, color=(71, 85, 105))

    doc.add_paragraph()
    disc = doc.add_paragraph()
    dr = disc.add_run(
        "Disclaimers: Isolation Forest scores are unsupervised outliers (“look here”), not fraud labels. "
        "Forecasts and decision bands guide planning envelopes and are not guarantees. "
        "Data are aggregate enrolment/update volumes only (no Aadhaar numbers or biometrics)."
    )
    set_run_font(dr, size=10, italic=True, color=(120, 53, 15))
    doc.add_page_break()

    i = 0
    if lines and lines[0].startswith("# "):
        i = 1

    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
                i += 1
                continue
            in_code = False
            code_text = "\n".join(code_buf)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            run = p.add_run(code_text if code_text else " ")
            set_run_font(run, size=8, font="Consolas", color=(30, 41, 59))
            pPr = p._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
            pPr.append(shd)
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped in ("---", "***", "___") or not stripped or stripped.startswith("<!--"):
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s\-:|]+\|$", lines[i + 1].strip()
        ):
            headers, rows, ni = parse_table(lines, i)
            if headers:
                add_table(doc, headers, rows)
                i = ni
                continue

        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m_img:
            alt, src = m_img.group(1), m_img.group(2)
            path = resolve_img(src)
            caption = None
            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if nxt.startswith("*") and nxt.endswith("*") and not nxt.startswith("**"):
                    caption = nxt.strip("*")
                    i += 1
            if path:
                add_image(doc, path, caption=caption or (alt if alt else None))
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"[FIGURE PLACEHOLDER] {alt or src}")
                set_run_font(r, size=9, italic=True, color=(37, 99, 235))
                if caption:
                    c = doc.add_paragraph()
                    cr = c.add_run(caption)
                    set_run_font(cr, size=9, italic=True, color=(71, 85, 105))
            i += 1
            continue

        if stripped.startswith(">"):
            qtext = stripped.lstrip("> ").strip()
            while i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                i += 1
                nxt = lines[i].strip().lstrip("> ").strip()
                if nxt:
                    qtext += " " + nxt
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(qtext)
            set_run_font(run, size=10, italic=True, color=(100, 116, 139))
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                item = re.sub(r"^[-*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_inline_markdown(p, item)
                i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_inline_markdown(p, item)
                i += 1
            continue

        p = doc.add_paragraph()
        content = stripped
        if content.startswith("*") and content.endswith("*") and not content.startswith("**"):
            content = content.strip("*")
            r = p.add_run(content)
            set_run_font(r, size=10, italic=True, color=(71, 85, 105))
            i += 1
            continue

        add_inline_markdown(p, content)
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
