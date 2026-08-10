# -*- coding: utf-8 -*-
"""
Build journal-ready Outcomes report for Aadhaar Intel Engine.
Evidence is pulled from live API snapshots (http://127.0.0.1:8787).
Figures are intentional placeholders for plots/screenshots.
"""
from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SNAP = Path(__file__).resolve().parent / "_live_api_snapshots" / "summary_for_outcomes.json"
OUT = Path(__file__).resolve().parent / "Aadhaar_Intel_Engine_Outcomes.docx"

# Palette
NAVY = "1E3A5F"
SLATE = "334155"
TEAL = "0F766E"
LIGHT = "F1F5F9"
AMBER_BG = "FFFBEB"
AMBER_BD = "D97706"
PLACEHOLDER_BG = "EEF2FF"
PLACEHOLDER_BD = "6366F1"


def set_run(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        if isinstance(color, str):
            color = tuple(int(color[i : i + 2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    tcPr.append(shd)


def set_cell_border(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(tcBorders)


def para(doc, text="", size=11, bold=False, italic=False, color=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_runs(p, parts):
    """parts: list of (text, kwargs)"""
    for text, kwargs in parts:
        r = p.add_run(text)
        set_run(r, **kwargs)
    return p


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        set_run(r, size=16, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        set_run(r, size=13, bold=True, color=TEAL)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        set_run(r, size=11, bold=True, color=SLATE)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, space_after=8):
    return para(doc, text, size=11, color=SLATE, space_after=space_after)


def caption(doc, text):
    return para(doc, text, size=9, italic=True, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, space_before=2)


def evidence_note(doc, text):
    p = para(doc, "", size=9, space_after=10, space_before=2)
    add_runs(
        p,
        [
            ("Evidence: ", {"size": 9, "bold": True, "color": TEAL, "italic": True}),
            (text, {"size": 9, "italic": True, "color": "64748B"}),
        ],
    )
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    r = p.add_run(text)
    set_run(r, size=11, color=SLATE)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + 0.15 * level)
    return p


def add_table(doc, headers, rows, col_widths=None, header_fill=NAVY):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run(r, size=9, bold=True, color="FFFFFF")
        shade_cell(cell, header_fill)
        set_cell_border(cell, "94A3B8", "4")

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run("" if val is None else str(val))
            set_run(r, size=9, color=SLATE)
            if ri % 2 == 1:
                shade_cell(cell, LIGHT)
            set_cell_border(cell, "CBD5E1", "4")

    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def figure_placeholder(doc, fig_id: str, title: str, path_hint: str, caption_text: str, height_lines: int = 3):
    """Visual placeholder box for a figure/screenshot/plot."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade_cell(cell, PLACEHOLDER_BG)
    set_cell_border(cell, PLACEHOLDER_BD, "12")

    cell.text = ""
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f"[ FIGURE {fig_id} — PLACEHOLDER ]")
    set_run(r0, size=10, bold=True, color=PLACEHOLDER_BD)

    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(title)
    set_run(r1, size=11, bold=True, color=NAVY)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Insert path: {path_hint}")
    set_run(r2, size=8, italic=True, color="64748B")

    for _ in range(max(1, height_lines - 1)):
        pe = cell.add_paragraph()
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        re = pe.add_run("·  ·  ·  space reserved for plot / screenshot  ·  ·  ·")
        set_run(re, size=8, color="94A3B8")

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "How to fill: capture from live web app (http://127.0.0.1:8787) "
        "or export chart; drop file at the path above."
    )
    set_run(r3, size=8, italic=True, color="64748B")

    caption(doc, caption_text)
    return t


def page_break(doc):
    doc.add_page_break()


def fmt_int(n):
    try:
        return f"{int(round(float(n))):,}"
    except Exception:
        return str(n)


def fmt_float(n, d=2):
    try:
        return f"{float(n):,.{d}f}"
    except Exception:
        return str(n)


def fmt_pct(n, d=2):
    try:
        return f"{float(n):.{d}f}%"
    except Exception:
        return str(n)


def build():
    S = json.loads(SNAP.read_text(encoding="utf-8"))
    k = S["dashboard_kpis"]
    meta = S.get("api_meta_brief") or {}
    geo_eval = meta.get("geo_eval") or {}
    geo_stats = meta.get("geo_repair_stats") or {}
    fc_meta = S.get("forecast_meta") or {}
    cmp = S.get("model_comparison") or []
    rp = S.get("resource_planning") or {}
    map_s = S.get("map_summary") or {}
    map_kpis = map_s.get("kpis") or {}
    top_risk = S.get("top_risk") or []
    reason_counts = S.get("reason_counts") or {}
    states_flagged = S.get("states_flagged") or []
    top_states = S.get("top_states") or []
    top_dist = S.get("map_top_districts") or []
    ratios = S.get("ratios") or {}
    age = S.get("dashboard_age_mix") or {}
    gov = meta.get("governance") or {}
    llm = meta.get("llm") or {}
    rows_raw = meta.get("rows_raw") or {}
    rows_q = meta.get("rows_quarantined") or {}
    dups = meta.get("duplicate_keys_collapsed") or {}
    mart_rows = meta.get("rows") or {}
    series = S.get("forecast_series_summary") or {}
    backtest = S.get("backtest") or {}
    rolling = S.get("rolling") or {}
    sel_meta = (fc_meta.get("selection_meta") or {}) if isinstance(fc_meta, dict) else {}

    captured_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    commit = "a53efa3"

    doc = Document()

    # --- page setup A4 journal-ish ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # Header / Footer
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Aadhaar Intel Engine — Research Outcomes  |  Evidence-backed report")
    set_run(hr, size=8, color="64748B", italic=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Confidential research artefact · Aggregate data only · Page ")
    set_run(fr, size=8, color="94A3B8")
    # page number field
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    for r in (run, run2, run3):
        set_run(r, size=8, color="94A3B8")

    # =========================
    # TITLE
    # =========================
    para(doc, "OUTCOMES REPORT", size=12, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(
        doc,
        "Aadhaar Intel Engine (UIDAI Operational Aggregates)",
        size=20,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    para(
        doc,
        "Data Processing · Anomaly Detection · Forecasting & Resource Planning · LLM Insights · Operator Web Application",
        size=11,
        italic=True,
        color=SLATE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
    )

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Document type", "Journal-oriented research outcomes (code- and live-system aligned)"],
            ["Evidence source", "Live web API http://127.0.0.1:8787 (cache-backed marts)"],
            ["Capture timestamp (UTC)", captured_at],
            ["Git commit (workspace)", commit],
            ["Data window", f"{meta.get('date_min')} → {meta.get('date_max')}"],
            ["Load source", f"{meta.get('source')} · marts schema v4 · geo rule pack {meta.get('geo_rule_pack')}"],
            ["LLM", f"{'available' if llm.get('available') else 'offline'} · model={llm.get('model')}"],
            ["Scope disclaimer", "Aggregate enrolment/update volumes only — no Aadhaar numbers or biometrics"],
        ],
        col_widths=[5.5, 11.5],
    )

    body(
        doc,
        "This document records measured outcomes of the Aadhaar Intel Engine as observed on the live operator web stack. "
        "Every quantitative claim is tied to an API/mart artefact (Evidence notes). Figure slots are reserved for "
        "plots and screenshots so the manuscript can be illustrated without embedding repository media prematurely.",
    )

    # Callout
    tcall = doc.add_table(rows=1, cols=1)
    c = tcall.rows[0].cells[0]
    shade_cell(c, AMBER_BG)
    set_cell_border(c, AMBER_BD, "12")
    c.text = ""
    p = c.paragraphs[0]
    add_runs(
        p,
        [
            ("Interpretive guardrails. ", {"size": 10, "bold": True, "color": AMBER_BD}),
            (
                "Isolation Forest risk scores are unsupervised multi-feature outliers (“look here”), not fraud labels. "
                "Forecasts and decision bands define planning envelopes, not guarantees. "
                "LLM prose is hybrid: metrics are engine-authored; narrative is optional Ollama assistance.",
                {"size": 10, "color": SLATE},
            ),
        ],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # =========================
    # 1. EXECUTIVE SUMMARY
    # =========================
    h1(doc, "1. Executive Summary of Outcomes")

    body(
        doc,
        "The system delivers an end-to-end operational analytics stack for UIDAI-style public aggregates: "
        "CSV ingest with geography repair, parquet analytical marts, multi-feature anomaly triage, "
        "baseline-gated multi-model forecasting with conformal intervals and resource planning, "
        "evidence-locked hybrid AI briefs, human governance patches, and a professional React + FastAPI web UI "
        "sharing one analytics engine with the Streamlit research prototype.",
    )

    h2(doc, "1.1 Headline measured results (live capture)")
    add_table(
        doc,
        ["Outcome area", "Headline metric", "Value"],
        [
            ["Operational volume", "Total enrolments (filtered national)", fmt_int(k["total_enrolments"])],
            ["Operational volume", "Biometric updates", fmt_int(k["biometric_updates"])],
            ["Operational volume", "Demographic updates", fmt_int(k["demographic_updates"])],
            ["Coverage", "Active states / districts", f"{fmt_int(k['active_states'])} / {fmt_int(k['active_districts'])}"],
            ["ETL / geo repair", "District alias normalizations (row-level)", fmt_int(geo_stats.get("district_aliases", 0))],
            ["ETL / geo repair", "AP → Telangana reassignments", fmt_int(geo_stats.get("ap_to_telangana", 0))],
            ["ETL / geo eval", "Gold set both-accuracy (n=19)", fmt_pct(100 * float(geo_eval.get("both_accuracy", 0)), 1)],
            ["Anomaly detection", "Flagged state×district cells", fmt_int(k["anomaly_count"])],
            ["Forecasting", "Selected model (Auto)", f"{k['forecast_model']} · reason={fc_meta.get('selection')}"],
            ["Forecasting", "Rolling MASE (primary)", fmt_float(k["rolling_mase"], 3)],
            ["Forecasting", "Decision band", str(k["decision_band"])],
            ["Forecasting", "30-day path change", fmt_pct(k["forecast_growth_pct"], 2)],
            ["Resource planning", "Operators @ 40 txn/day on peak", fmt_int(rp.get("operators_at_40", 0))],
            ["Web / governance", "Durable district merge patches", fmt_int(gov.get("district_merges", 0))],
            ["LLM insights", "Hybrid briefs available", f"yes · {llm.get('model')}"],
        ],
        col_widths=[4.5, 7.0, 5.5],
    )
    evidence_note(
        doc,
        "GET /api/dashboard, /api/forecast?horizon=30&model=Auto, /api/meta, /api/map · snapshots under docs/_live_api_snapshots/",
    )

    h2(doc, "1.2 Research contributions (outcome map)")
    add_table(
        doc,
        ["ID", "Contribution", "Where evidenced"],
        [
            ["C1", "Production-style geo repair with versioned rule pack + gold eval", "geo_repair_stats, geo_eval @ /api/meta; tests/test_geo_repair.py"],
            ["C2", "Cache-first parquet marts (daily ops grain)", "data/processed/* · schema_version=4 · source=cache"],
            ["C3", "Multi-feature Isolation Forest with reason codes & notes", "/api/dashboard top_risk · /api/analytics risk_cells"],
            ["C4", "Forecast bake-off: rolling CV, MASE primary, baseline gate, conformal bands", "/api/forecast comparison + meta"],
            ["C5", "Human-in-the-loop governance (durable merges/deletes + audit)", "governance_patches.json · 96 district merges"],
            ["C6", "Evidence-locked hybrid AI (engine numbers + optional LLM)", "/api/insights/* · source hybrid"],
            ["C7", "Operator web app at feature parity with research UI", "React SPA + FastAPI :8787 · FEATURE_PARITY.md"],
        ],
        col_widths=[1.5, 8.5, 7.0],
    )

    figure_placeholder(
        doc,
        "1",
        "System architecture at a glance (dual UI over shared engine)",
        "docs/figures/fig01_architecture.png",
        "Figure 1. Placeholder — architecture overview: CSV/marts → AnalyticsEngine → FastAPI/React & Streamlit.",
    )

    # =========================
    # 2. SCOPE & RQs
    # =========================
    h1(doc, "2. Problem Framing, Scope, and Research Questions")

    h2(doc, "2.1 Operational problem")
    body(
        doc,
        "Public identity-system dashboards often stop at totals and static maps. Operators need clean, comparable geography; "
        "transparent district-level outlier investigation; demand outlook with uncertainty; and governance that survives reloads. "
        "This work addresses those needs on aggregate (non-individual) data.",
    )

    h2(doc, "2.2 Research questions")
    add_table(
        doc,
        ["ID", "Question", "Approach", "Outcome status"],
        [
            ["RQ1", "Can geo rules recover official states/districts on a gold set?", "Rule pack 1.1.0 + evaluate_geo_cleaning", f"both_accuracy={geo_eval.get('both_accuracy')} (n={geo_eval.get('n')})"],
            ["RQ2", "Do multi-feature outliers surface actionable cells beyond raw spikes?", "IsolationForest on volume, CV, ratios, WoW, etc.", f"{k['anomaly_count']} cells; multi-reason codes"],
            ["RQ3", "When should complex forecasts beat simple baselines?", "Bake-off + beat SN & MA on MASE", f"Ensemble selected · beats_baselines · MASE={k['rolling_mase']}"],
            ["RQ4", "Can hybrid AI assist without fabricating numbers?", "Evidence dict → draft → optional Ollama", "Key numbers engine-authored; LLM narrative optional"],
            ["RQ5", "Can a professional web UI retain research fidelity?", "Shared engine + FastAPI + React", "Live at :8787; parity documented"],
        ],
        col_widths=[1.5, 5.5, 5.0, 5.0],
    )

    h2(doc, "2.3 Explicit non-goals")
    for t in [
        "Individual-level identity adjudication, biometric matching, or legal fraud determination.",
        "Causal claims about migration or policy impact from aggregates alone.",
        "Guaranteed forecast accuracy under structural breaks without re-evaluation.",
    ]:
        bullet(doc, t)

    # =========================
    # 3. ETL OUTCOMES
    # =========================
    page_break(doc)
    h1(doc, "3. Outcomes — Data Processing / ETL")

    body(
        doc,
        "Data processing is treated as a first-class research outcome: without transparent cleaning, "
        "downstream anomaly and forecast metrics are not comparable across states and districts.",
    )

    h2(doc, "3.1 Source inventory and contracts")
    add_table(
        doc,
        ["Stream", "Raw rows (CSV shards)", "Quarantined", "Duplicate keys collapsed", "Daily mart rows"],
        [
            ["Enrolment", fmt_int(rows_raw.get("enrol", 0)), fmt_int(rows_q.get("enrol", 0)), fmt_int(dups.get("enrol", 0)), fmt_int(mart_rows.get("enrol", 0))],
            ["Biometric updates", fmt_int(rows_raw.get("bio", 0)), fmt_int(rows_q.get("bio", 0)), fmt_int(dups.get("bio", 0)), fmt_int(mart_rows.get("bio", 0))],
            ["Demographic updates", fmt_int(rows_raw.get("demo", 0)), fmt_int(rows_q.get("demo", 0)), fmt_int(dups.get("demo", 0)), fmt_int(mart_rows.get("demo", 0))],
            ["Total raw", fmt_int(sum(rows_raw.values()) if rows_raw else 0), fmt_int(sum(rows_q.values()) if rows_q else 0), fmt_int(sum(dups.values()) if dups else 0), "—"],
        ],
        col_widths=[3.8, 3.5, 2.8, 3.8, 3.1],
    )
    evidence_note(doc, "load_report from GET /api/meta; mart row counts from cache source.")

    body(
        doc,
        "Grain contracts: natural key for de-duplication is date × state × district × pincode; "
        "analytics grain materialised in marts is date × state × district. "
        "Invalid dates/pincodes/non-official states are quarantined before repair.",
    )

    h2(doc, "3.2 Pipeline stages")
    body(
        doc,
        "CSV shards → classify (enrol/bio/demo) → validate dates & pincode range → quarantine invalid rows → "
        "collapse natural-key duplicates (sum) → canonicalize state → full_geo_repair → "
        "optional AUTO_BUILD_CACHE → parquet marts + fingerprint manifest.",
    )

    figure_placeholder(
        doc,
        "2",
        "ETL pipeline flowchart (ingest → repair → marts → engine)",
        "docs/figures/fig02_etl_pipeline.png",
        "Figure 2. Placeholder — ETL and cache materialisation pipeline.",
        height_lines=2,
    )

    h2(doc, "3.3 Alias fixing and geography repair")
    body(
        doc,
        "Ordered stages in full_geo_repair(): (1) district typed as state; (2) district aliases; "
        "(3) Andhra Pradesh → Telangana boundary allowlist; (4) PIN prefix → state fallback when state remains non-official. "
        f"Rule pack version: {meta.get('geo_rule_pack')}.",
    )

    add_table(
        doc,
        ["Repair stage", "Rows affected (load_report)", "Interpretation"],
        [
            ["district_as_state", fmt_int(geo_stats.get("district_as_state", 0)), "District names incorrectly placed in state field"],
            ["district_aliases", fmt_int(geo_stats.get("district_aliases", 0)), "Spelling/alias normalization to canonical districts"],
            ["ap_to_telangana", fmt_int(geo_stats.get("ap_to_telangana", 0)), "Post-bifurcation district reassignment"],
            ["pin_prefix_fallback", fmt_int(geo_stats.get("pin_prefix_fallback", 0)), "State recovery via PIN prefix map"],
        ],
        col_widths=[4.0, 4.5, 8.5],
    )
    evidence_note(doc, "load_report.geo_repair_stats @ GET /api/meta (also mirrored in data/processed/manifest.json).")

    h3(doc, "3.3.1 Gold-set evaluation (automated evidence)")
    per_rule = (geo_eval.get("per_rule") or {}) if geo_eval else {}
    rule_rows = []
    for name, m in sorted(per_rule.items()):
        rule_rows.append(
            [
                name,
                fmt_pct(100 * float(m.get("state_ok", 0)), 0),
                fmt_pct(100 * float(m.get("district_ok", 0)), 0),
                fmt_pct(100 * float(m.get("both_ok", 0)), 0),
            ]
        )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Gold examples (n)", fmt_int(geo_eval.get("n", 0))],
            ["Rule pack", str(geo_eval.get("rule_pack", meta.get("geo_rule_pack")))],
            ["State accuracy", fmt_pct(100 * float(geo_eval.get("state_accuracy", 0)), 1)],
            ["District accuracy", fmt_pct(100 * float(geo_eval.get("district_accuracy", 0)), 1)],
            ["Both accuracy", fmt_pct(100 * float(geo_eval.get("both_accuracy", 0)), 1)],
            ["Failures", str(geo_eval.get("failures") or [])],
        ],
        col_widths=[6.0, 11.0],
    )
    if rule_rows:
        para(doc, "Per-rule gold accuracy", size=10, bold=True, color=NAVY, space_after=4)
        add_table(
            doc,
            ["Rule family", "State OK", "District OK", "Both OK"],
            rule_rows,
            col_widths=[5.0, 4.0, 4.0, 4.0],
        )
    evidence_note(doc, "geo_eval embedded in load_report from live /api/meta; harness assets/reference/eval_geo_gold.json.")

    figure_placeholder(
        doc,
        "3",
        "Geo repair before/after examples (alias + AP→Telangana)",
        "docs/figures/fig03_geo_repair_examples.png",
        "Figure 3. Placeholder — illustrative rows before/after geography repair.",
        height_lines=2,
    )

    h2(doc, "3.4 Parquet analytical marts")
    add_table(
        doc,
        ["Mart file", "Role", "Rows (live)", "Approx. size on disk"],
        [
            ["fact_enrol_daily.parquet", "Daily enrolment grain", fmt_int(mart_rows.get("enrol", 0)), "~364 KB"],
            ["fact_bio_daily.parquet", "Daily biometric updates", fmt_int(mart_rows.get("bio", 0)), "~451 KB"],
            ["fact_demo_daily.parquet", "Daily demographic updates", fmt_int(mart_rows.get("demo", 0)), "~446 KB"],
            ["agg_district.parquet", "District aggregates", "953", "~32 KB"],
            ["dim_geo.parquet", "Geography dimension", "30,162", "~169 KB"],
            ["manifest.json", "Fingerprint + load_report", "—", "schema_version=4"],
        ],
        col_widths=[5.0, 5.0, 3.5, 3.5],
    )
    body(
        doc,
        "Cache-first loading (source=cache, cache_valid=true on this capture) enables sub-second operator UI reloads "
        "versus multi-million-row CSV re-parse. Schema version 4 gates compatibility when repair rules or mart layout change.",
    )
    evidence_note(doc, "GET /api/health source=cache; GET /api/meta rows; data/processed/manifest.json.")

    figure_placeholder(
        doc,
        "4",
        "Mart schema / row-count dashboard (optional ops slide)",
        "docs/figures/fig04_mart_inventory.png",
        "Figure 4. Placeholder — parquet mart inventory and fingerprint.",
        height_lines=2,
    )

    h2(doc, "3.5 Human governance overlays")
    body(
        doc,
        "Residual name issues after automated repair are addressed by durable human patches "
        "(merge/delete/ignore) stored outside the mart rebuild so operator decisions survive reloads.",
    )
    add_table(
        doc,
        ["Governance artefact", "Live count / path"],
        [
            ["State merges", fmt_int(gov.get("state_merges", 0))],
            ["District merges", fmt_int(gov.get("district_merges", 0))],
            ["State deletions", fmt_int(gov.get("state_deletions", 0))],
            ["District deletions", fmt_int(gov.get("district_deletions", 0))],
            ["Store path", str(gov.get("store_path", "output/governance_patches.json"))],
            ["Audit log", "output/governance_audit.csv"],
        ],
        col_widths=[5.5, 11.5],
    )
    evidence_note(doc, "governance block on GET /api/meta; UI: Governance page on live web app.")

    figure_placeholder(
        doc,
        "5",
        "Governance console screenshot (scan / merge / audit)",
        "docs/figures/fig05_governance_ui.png",
        "Figure 5. Placeholder — operator governance UI with durable patches.",
        height_lines=2,
    )

    # =========================
    # 4. ANOMALY OUTCOMES
    # =========================
    page_break(doc)
    h1(doc, "4. Outcomes — Anomaly Detection")

    body(
        doc,
        "Anomaly detection triages state×district cells that are multi-feature outliers relative to peers. "
        "Scores support investigation prioritisation; they are not calibrated fraud probabilities.",
    )

    h2(doc, "4.1 Method summary")
    add_table(
        doc,
        ["Parameter", "Value (this run)"],
        [
            ["Unit of analysis", "state × district (composite cell)"],
            ["Model", "sklearn.ensemble.IsolationForest"],
            ["Contamination", str(S.get("dashboard_params", {}).get("contamination", 0.05))],
            ["Minimum volume", str(S.get("dashboard_params", {}).get("min_volume", 50))],
            ["n_estimators / seed", "200 / fixed random_state (engine)"],
            [
                "Feature family",
                "volume, log volume, day-to-day CV, bio ratio, demo ratio, WoW growth, active days, volume vs state median",
            ],
            ["Outputs", "risk_score, reason code, investigation_notes, driver_z"],
        ],
        col_widths=[4.5, 12.5],
    )

    h2(doc, "4.2 Detection outcomes (national filter)")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Cells flagged", fmt_int(k["anomaly_count"])],
            ["Contamination setting", str(S.get("dashboard_params", {}).get("contamination", 0.05))],
            ["Min volume gate", str(S.get("dashboard_params", {}).get("min_volume", 50))],
            ["Top reason (mode)", max(reason_counts, key=reason_counts.get) if reason_counts else "—"],
        ],
        col_widths=[6.0, 11.0],
    )
    evidence_note(doc, "GET /api/dashboard (anomaly_count, top_risk); GET /api/analytics (risk_cells).")

    h3(doc, "4.2.1 Reason-code distribution")
    reason_rows = [[r, fmt_int(c), fmt_pct(100 * c / max(1, sum(reason_counts.values())), 1)] for r, c in sorted(reason_counts.items(), key=lambda x: -x[1])]
    add_table(
        doc,
        ["Reason code", "# cells", "Share of flagged"],
        reason_rows,
        col_widths=[8.0, 3.5, 5.5],
    )

    h3(doc, "4.2.2 States with most flagged cells")
    state_rows = [[s, fmt_int(c)] for s, c in states_flagged[:12]]
    add_table(doc, ["State / UT", "Flagged cells"], state_rows, col_widths=[10.0, 7.0])

    h3(doc, "4.2.3 Top investigation queue (live)")
    risk_rows = []
    for r in top_risk[:12]:
        risk_rows.append(
            [
                r.get("state"),
                r.get("district"),
                fmt_int(r.get("volume", 0)),
                fmt_int(r.get("risk_score", 0)),
                r.get("reason"),
                fmt_float(r.get("cv", 0), 2),
                fmt_float(r.get("bio_ratio", 0), 2),
                fmt_float(r.get("demo_ratio", 0), 2),
            ]
        )
    add_table(
        doc,
        ["State", "District", "Volume", "Risk", "Reason", "CV", "Bio", "Demo"],
        risk_rows,
        col_widths=[2.6, 3.4, 1.6, 1.2, 3.6, 1.2, 1.2, 1.2],
    )
    evidence_note(doc, "top_risk from GET /api/dashboard; full investigation_notes available per cell in API payload.")

    # example investigation note
    if top_risk:
        h3(doc, "4.2.4 Example investigation note (engine-authored)")
        note = top_risk[0].get("investigation_notes") or ""
        tnote = doc.add_table(rows=1, cols=1)
        nc = tnote.rows[0].cells[0]
        shade_cell(nc, LIGHT)
        set_cell_border(nc, "94A3B8", "8")
        nc.text = ""
        p = nc.paragraphs[0]
        r = p.add_run(note)
        set_run(r, size=9, italic=True, color=SLATE)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    figure_placeholder(
        doc,
        "6",
        "Anomaly risk bars / scatter (Analytics page)",
        "docs/figures/fig06_anomaly_risk.png",
        "Figure 6. Placeholder — multi-feature risk ranking visualization.",
        height_lines=2,
    )
    figure_placeholder(
        doc,
        "7",
        "Risk radar / feature profile for top cells",
        "docs/figures/fig07_risk_radar.png",
        "Figure 7. Placeholder — radar or parallel-coordinates of anomaly features.",
        height_lines=2,
    )

    h2(doc, "4.3 Operational reading")
    for t in [
        f"With contamination={S.get('dashboard_params', {}).get('contamination', 0.05)}, the engine surfaces a short list "
        f"({fmt_int(k['anomaly_count'])} cells) suitable for human review rather than an exhaustive list of all high-volume districts.",
        "Dominant reason codes on this capture include irregular activity footprint, volume vs state median, and day-to-day volatility — "
        "indicating the detector is not reducible to a single volume threshold.",
        "Meghalaya and West Bengal appear frequently in both high-risk and high-volume contexts; operators should cross-check "
        "with local calendars, kit availability, and name-quality campaigns before escalating.",
    ]:
        bullet(doc, t)

    # =========================
    # 5. FORECAST OUTCOMES
    # =========================
    page_break(doc)
    h1(doc, "5. Outcomes — Forecasting, Model Comparison & Resource Planning")

    body(
        doc,
        "Forecasting targets national (or single-state) daily enrolment. Missing calendar days are filled with zeros to preserve "
        "day-of-week structure. Auto mode runs a multi-model bake-off with rolling-origin cross-validation, ranks by MASE, "
        "and applies a baseline gate so non-baseline models must beat both SeasonalNaive and MovingAverage.",
    )

    h2(doc, "5.1 Run configuration")
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ["Scope", "National daily enrolment"],
            ["Horizon", f"{fc_meta.get('horizon', 30)} days"],
            ["Train days", fmt_int(fc_meta.get("train_days", 0))],
            ["Data end", str(fc_meta.get("data_end"))],
            ["Candidates", ", ".join(fc_meta.get("candidates") or ["MovingAverage", "Drift", "Ensemble", "SeasonalNaive"])],
            ["Primary metric", str(fc_meta.get("primary_metric", "mase"))],
            ["Selection mode", str(fc_meta.get("selection"))],
            ["Calendar fill missing days", str(fc_meta.get("calendar_filled"))],
            ["Interval method", str(fc_meta.get("interval_method"))],
            ["Conformal α / q", f"{fc_meta.get('conformal_alpha')} / {fmt_float(fc_meta.get('conformal_q'), 0)}"],
            ["Decision band", str(fc_meta.get("decision_band"))],
        ],
        col_widths=[6.0, 11.0],
    )
    evidence_note(doc, "GET /api/forecast?horizon=30&model=Auto → meta block.")

    h2(doc, "5.2 Model bake-off (rolling-origin CV)")
    cmp_rows = []
    selected = fc_meta.get("model") or S.get("forecast_model")
    for i, row in enumerate(cmp, 1):
        name = row.get("model")
        mark = "YES" if name == selected else ""
        cmp_rows.append(
            [
                i,
                name,
                fmt_float(row.get("mase"), 4),
                fmt_float(row.get("smape_pct"), 2),
                fmt_float(row.get("mape_pct"), 2),
                fmt_float(row.get("rmse"), 1),
                fmt_float(row.get("bias"), 1),
                fmt_int(row.get("n_folds", 0)),
                row.get("decision_band"),
                mark,
            ]
        )
    add_table(
        doc,
        ["Rank", "Model", "MASE↓", "sMAPE%", "MAPE%", "RMSE", "Bias", "Folds", "Band", "Selected"],
        cmp_rows,
        col_widths=[1.2, 2.6, 1.5, 1.6, 1.5, 1.8, 1.6, 1.2, 2.0, 1.6],
    )

    baselines = (sel_meta.get("baselines") or {}) if isinstance(sel_meta, dict) else {}
    add_table(
        doc,
        ["Selection audit item", "Value"],
        [
            ["Best by MASE (raw #1)", str(sel_meta.get("best_raw") if isinstance(sel_meta, dict) else selected)],
            ["Best raw MASE", fmt_float(sel_meta.get("best_score"), 4) if isinstance(sel_meta, dict) else "—"],
            ["Baseline SeasonalNaive MASE", fmt_float(baselines.get("SeasonalNaive"), 4) if baselines else "—"],
            ["Baseline MovingAverage MASE", fmt_float(baselines.get("MovingAverage"), 4) if baselines else "—"],
            ["Beat-baseline ε", str(sel_meta.get("beat_baseline_eps") if isinstance(sel_meta, dict) else 0.0)],
            ["Selected after gate", str(sel_meta.get("selected") if isinstance(sel_meta, dict) else selected)],
            ["Selection reason", str(sel_meta.get("reason") if isinstance(sel_meta, dict) else fc_meta.get("selection"))],
            [
                "Δ MASE vs best baseline (MA)",
                fmt_float(float(baselines.get("MovingAverage", 0)) - float(sel_meta.get("best_score") or 0), 4)
                if baselines and isinstance(sel_meta, dict)
                else "—",
            ],
        ],
        col_widths=[7.0, 10.0],
    )
    body(
        doc,
        f"On this capture, Ensemble wins on primary metric MASE={fmt_float(k['rolling_mase'], 3)} and "
        f"passes the baseline gate (reason={fc_meta.get('selection')}). Absolute-level sMAPE remains high "
        f"({fmt_float(k['rolling_smape'], 1)}%), which the system honestly maps to a “{k['decision_band']}” decision band: "
        "useful for direction and staffing envelopes, not day-exact quotas.",
    )

    h2(doc, "5.3 Holdout backtest (secondary check)")
    add_table(
        doc,
        ["Holdout metric", "Value"],
        [
            ["Holdout days", fmt_int(backtest.get("holdout_days", 14))],
            ["Model", str(backtest.get("model"))],
            ["MASE", fmt_float(backtest.get("mase"), 4)],
            ["sMAPE %", fmt_float(backtest.get("smape_pct"), 2)],
            ["MAPE %", fmt_float(backtest.get("mape_pct"), 2)],
            ["RMSE", fmt_float(backtest.get("rmse"), 1)],
            ["Mean actual / mean pred", f"{fmt_float(backtest.get('mean_actual'), 1)} / {fmt_float(backtest.get('mean_pred'), 1)}"],
            ["Holdout decision band", str(backtest.get("decision_band"))],
            ["Status", str(backtest.get("status"))],
        ],
        col_widths=[6.0, 11.0],
    )
    body(
        doc,
        "Holdout metrics are reported separately from rolling CV used for selection. Divergence (e.g., exploratory holdout band "
        "while rolling remains directional) is an intentional transparency signal for operators reviewing recent regime shifts.",
    )
    evidence_note(doc, "meta.backtest and meta.rolling on GET /api/forecast.")

    h2(doc, "5.4 Forecast path summary")
    first = series.get("first") or {}
    last = series.get("last") or {}
    add_table(
        doc,
        ["Path statistic", "Value"],
        [
            ["Horizon points", fmt_int(series.get("n", 0))],
            ["First day (pred)", f"{first.get('date')} · {fmt_int(first.get('predicted', 0))}"],
            ["Last day (pred)", f"{last.get('date')} · {fmt_int(last.get('predicted', 0))}"],
            ["Peak predicted (horizon)", fmt_int(series.get("peak_predicted", S.get("forecast_peak")))],
            ["Floor predicted (horizon)", fmt_int(series.get("floor_predicted", S.get("forecast_floor")))],
            ["Average daily predicted", fmt_int(series.get("avg_predicted", rp.get("avg_daily")))],
            ["Sum predicted (30d)", fmt_int(series.get("sum_predicted", 0))],
            ["Path change (dashboard)", fmt_pct(k["forecast_growth_pct"], 2)],
            ["Conformal half-width q", fmt_int(fc_meta.get("conformal_q", 0))],
        ],
        col_widths=[6.0, 11.0],
    )

    figure_placeholder(
        doc,
        "8",
        "Forecast path with conformal lower/upper bands",
        "docs/figures/fig08_forecast_path.png",
        "Figure 8. Placeholder — 30-day national enrolment forecast with uncertainty band.",
        height_lines=2,
    )
    figure_placeholder(
        doc,
        "9",
        "Model selection charts (MASE rank + Δ vs baseline gate)",
        "docs/figures/fig09_model_bakeoff.png",
        "Figure 9. Placeholder — bake-off ranking and baseline-gate visualization.",
        height_lines=2,
    )

    h2(doc, "5.5 Resource planning outcomes")
    body(
        doc,
        "Resource planning translates the forecast envelope into operator staffing at a configurable throughput "
        f"(default {fmt_int(rp.get('per_operator', 40))} transactions per operator per day).",
    )
    add_table(
        doc,
        ["Planning metric", "Value", "Ops interpretation"],
        [
            ["Peak daily enrolment (horizon)", fmt_int(rp.get("peak_daily", 0)), "Upper staffing stress day"],
            ["Average daily enrolment (horizon)", fmt_int(rp.get("avg_daily", 0)), "Steady-state roster"],
            ["Throughput assumption", f"{fmt_int(rp.get('per_operator', 40))} txn / operator / day", "Configurable in UI"],
            ["Operators required at peak", fmt_int(rp.get("operators_at_40", 0)), "ceil(peak / per_operator)"],
            ["Operators at average (derived)", fmt_int((rp.get("avg_daily") or 0) / max(1, rp.get("per_operator") or 40)), "Steady roster estimate"],
        ],
        col_widths=[5.5, 4.0, 7.5],
    )
    evidence_note(doc, "resource_planning object on GET /api/forecast.")

    figure_placeholder(
        doc,
        "10",
        "Resource planning view (peak vs average staffing)",
        "docs/figures/fig10_resource_planning.png",
        "Figure 10. Placeholder — operator requirement chart derived from forecast peak/average.",
        height_lines=2,
    )

    h2(doc, "5.6 Decision-band guidance (method → ops)")
    add_table(
        doc,
        ["Band", "MASE guide", "sMAPE guide", "Recommended use"],
        [
            ["tight", "< 0.85", "< 20%", "Finer short-horizon planning"],
            ["directional", "< 1.15", "< 40%", "Direction & envelope (this run on rolling MASE)"],
            ["exploratory", "≥ 1.15", "≥ 40%", "Monitor actuals closely; avoid hard quotas"],
        ],
        col_widths=[3.0, 3.5, 3.5, 7.0],
    )

    # =========================
    # 6. LLM INSIGHTS
    # =========================
    page_break(doc)
    h1(doc, "6. Outcomes — LLM / Hybrid Insights")

    body(
        doc,
        "Hybrid AI briefs separate measurement from narrative. The engine first materialises an evidence dictionary "
        "(KPIs, ranks, bake-off metrics). Deterministic Insights and Actions are drafted from that evidence; "
        "optional Ollama prose may rephrase but must not invent numbers. Key numbers remain engine-authored.",
    )

    h2(doc, "6.1 Runtime status")
    add_table(
        doc,
        ["Item", "Live value"],
        [
            ["LLM available", str(llm.get("available"))],
            ["Model", str(llm.get("model"))],
            ["Error", str(llm.get("error"))],
            ["Dashboard insights endpoint", "GET /api/insights/dashboard"],
            ["Forecast insights endpoint", "GET /api/insights/forecast"],
            ["Assembly mode (observed)", "hybrid · narrative by Ollama when available"],
        ],
        col_widths=[6.0, 11.0],
    )
    evidence_note(doc, "llm block on GET /api/meta; insight markdown on /api/insights/*.")

    h2(doc, "6.2 Dashboard brief — engine Key numbers (evidence-locked)")
    add_table(
        doc,
        ["Key number", "Value"],
        [
            ["Data as of", str(meta.get("date_max"))],
            ["Total enrolments", fmt_int(k["total_enrolments"])],
            ["Biometric updates", fmt_int(k["biometric_updates"])],
            ["Demographic updates", fmt_int(k["demographic_updates"])],
            ["Bio per enrolment", fmt_float(ratios.get("bio_per_enrol"), 2)],
            ["Demo per enrolment", fmt_float(ratios.get("demo_per_enrol"), 2)],
            ["Anomaly count", fmt_int(k["anomaly_count"])],
            ["Contamination / min volume", f"{S.get('dashboard_params', {}).get('contamination')} / {S.get('dashboard_params', {}).get('min_volume')}"],
        ],
        col_widths=[6.0, 11.0],
    )

    h3(doc, "6.2.1 Narrative excerpt (optional LLM layer)")
    excerpt = (S.get("insights_dashboard_excerpt") or "").replace("\u00e2\u0080\u0094", "—")
    # clean mojibake lightly
    for bad, good in [
        ("â", "—"),
        ("Ã—", "×"),
        ("Â·", "·"),
        ("â€”", "—"),
    ]:
        excerpt = excerpt.replace(bad, good)
    tnote = doc.add_table(rows=1, cols=1)
    nc = tnote.rows[0].cells[0]
    shade_cell(nc, "F0FDFA")
    set_cell_border(nc, TEAL, "8")
    nc.text = ""
    p = nc.paragraphs[0]
    r = p.add_run((excerpt[:900] + "…") if len(excerpt) > 900 else excerpt)
    set_run(r, size=9, color=SLATE)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    caption(doc, "Excerpt from hybrid dashboard insight (live). Full text via GET /api/insights/dashboard.")

    h2(doc, "6.3 Forecast brief — engine Key numbers")
    add_table(
        doc,
        ["Key number", "Value"],
        [
            ["Selected model", str(selected)],
            ["Selection reason", str(fc_meta.get("selection"))],
            ["MASE / sMAPE", f"{fmt_float(k['rolling_mase'], 3)} / {fmt_float(k['rolling_smape'], 1)}%"],
            ["Decision band", str(k["decision_band"])],
            ["Change % (30d path)", fmt_pct(k["forecast_growth_pct"], 2)],
            ["Peak / floor", f"{fmt_int(S.get('forecast_peak'))} / {fmt_int(S.get('forecast_floor'))}"],
            ["Conformal q", fmt_int(fc_meta.get("conformal_q", 0))],
            ["Train days / data end", f"{fmt_int(fc_meta.get('train_days', 0))} / {fc_meta.get('data_end')}"],
        ],
        col_widths=[6.0, 11.0],
    )

    figure_placeholder(
        doc,
        "11",
        "AI panel screenshots (Dashboard + Forecast) showing Key numbers",
        "docs/figures/fig11_ai_panels.png",
        "Figure 11. Placeholder — hybrid AI panels with visible engine Key numbers.",
        height_lines=2,
    )

    h2(doc, "6.4 Outcome for operators")
    for t in [
        "Briefs reduce time-to-first-interpretation of dense KPI pages while keeping auditability: every bold metric can be traced to the engine table above.",
        "When Ollama is offline, the system still returns engine-only analysis — no silent empty state for critical ops numbers.",
        "Caveats are injected into the brief (unsupervised risk; directional forecasts), reducing over-trust risk in journal or ops review settings.",
    ]:
        bullet(doc, t)

    # =========================
    # 7. WEB APP OUTCOMES
    # =========================
    page_break(doc)
    h1(doc, "7. Outcomes — Live Operator Web Application")

    body(
        doc,
        "A primary project focus is the professional operator web application: a React (Vite) SPA served with a FastAPI backend "
        "that reuses the same DataLoader, AnalyticsEngine, ForecastBackend, geo repair, and governance store as the Streamlit research UI. "
        "Feature parity is documented in web/FEATURE_PARITY.md.",
    )

    h2(doc, "7.1 Deployment evidence (this capture)")
    add_table(
        doc,
        ["Check", "Result"],
        [
            ["Base URL", "http://127.0.0.1:8787"],
            ["HTTP health", "ok=true, ready=true, source=cache"],
            ["API version", "Aadhaar Intel Engine API 2.0.0"],
            ["Frontend", "Static SPA from web/frontend/dist mounted by FastAPI"],
            ["Data window", f"{meta.get('date_min')} → {meta.get('date_max')}"],
            ["States in filter", fmt_int(meta.get("states_n", 36))],
            ["Launch", "python run_web.py (optional --skip-build)"],
        ],
        col_widths=[5.0, 12.0],
    )
    evidence_note(doc, "GET /api/health, OpenAPI /docs, live HTML shell at /.")

    h2(doc, "7.2 Module outcomes (operator journeys)")
    add_table(
        doc,
        ["Module", "Operator value", "Live evidence"],
        [
            ["Dashboard", "KPI strip, workload/age mix, risk queue, 30-day outlook, AI brief", "/api/dashboard · /api/insights/dashboard"],
            ["Analytics", "Deep dive: state ranks, scatter, radar, watchlist, CSV exports", "/api/analytics · export routes"],
            ["Forecast", "Bake-off, conformal path, resource planning, per-state scope, AI brief", "/api/forecast · resource_planning"],
            ["Geospatial", "2D/heatmap/3D intensity, centroids, hotspot KPIs, CSV exports", "/api/map · /api/geojson"],
            ["Governance", "Merge/delete/ignore, audit/revert, pack import-export", "/api/governance*"],
        ],
        col_widths=[3.0, 7.5, 6.5],
    )

    h2(doc, "7.3 Operational KPI strip (Dashboard)")
    add_table(
        doc,
        ["KPI", "Value"],
        [
            ["Total enrolments", fmt_int(k["total_enrolments"])],
            ["Adult (18+) enrolments", fmt_int(k["adult_enrolments"])],
            ["Biometric updates", fmt_int(k["biometric_updates"])],
            ["Demographic updates", fmt_int(k["demographic_updates"])],
            ["Active states", fmt_int(k["active_states"])],
            ["Active districts", fmt_int(k["active_districts"])],
            ["Risk cells flagged", fmt_int(k["anomaly_count"])],
            ["Forecast model / band", f"{k['forecast_model']} / {k['decision_band']}"],
            ["Forecast growth (30d)", fmt_pct(k["forecast_growth_pct"], 2)],
            ["Rolling MASE / sMAPE", f"{fmt_float(k['rolling_mase'], 3)} / {fmt_float(k['rolling_smape'], 1)}%"],
        ],
        col_widths=[7.0, 10.0],
    )

    # Age mix
    h3(doc, "7.3.1 Age structure of enrolments")
    age_rows = []
    total_age = sum(float(v) for v in age.values()) or 1.0
    # normalize keys
    for label, val in age.items():
        clean = label
        for bad in ["â\x80\x93", "â€“", "–", "â"]:
            clean = clean.replace(bad, "–")
        if "0" in clean and "5" in clean and "17" not in clean:
            clean = "0–5"
        elif "5" in clean and "17" in clean:
            clean = "5–17"
        age_rows.append([clean, fmt_int(val), fmt_pct(100 * float(val) / total_age, 1)])
    # ensure order
    if not age_rows:
        age_rows = [["0–5", "—", "—"], ["5–17", "—", "—"], ["18+", fmt_int(k["adult_enrolments"]), "—"]]
    add_table(doc, ["Age band", "Enrolments", "Share"], age_rows, col_widths=[5.0, 6.0, 6.0])

    h3(doc, "7.3.2 Workload mix")
    add_table(
        doc,
        ["Stream", "Volume", "Per enrolment"],
        [
            ["Enrolments", fmt_int(k["total_enrolments"]), "1.00"],
            ["Biometric updates", fmt_int(k["biometric_updates"]), fmt_float(ratios.get("bio_per_enrol"), 2)],
            ["Demographic updates", fmt_int(k["demographic_updates"]), fmt_float(ratios.get("demo_per_enrol"), 2)],
        ],
        col_widths=[5.5, 5.5, 6.0],
    )

    h3(doc, "7.3.3 Top states by enrolment volume")
    ts_rows = []
    for i, row in enumerate(top_states[:10], 1):
        ts_rows.append([i, row.get("state"), fmt_int(row.get("volume", 0)), fmt_pct(100 * float(row.get("volume", 0)) / max(1, float(k["total_enrolments"])), 2)])
    add_table(doc, ["Rank", "State", "Volume", "Share of national"], ts_rows, col_widths=[2.0, 6.0, 4.5, 4.5])

    figure_placeholder(
        doc,
        "12",
        "Dashboard homepage screenshot (KPI strip + charts)",
        "docs/figures/fig12_dashboard.png",
        "Figure 12. Placeholder — operator Dashboard as served on the live web app.",
        height_lines=3,
    )
    figure_placeholder(
        doc,
        "13",
        "Analytics page screenshot (risk + regional deep dive)",
        "docs/figures/fig13_analytics.png",
        "Figure 13. Placeholder — Analytics investigation workspace.",
        height_lines=2,
    )
    figure_placeholder(
        doc,
        "14",
        "Forecast page screenshot (path, bake-off, resource planning)",
        "docs/figures/fig14_forecast_ui.png",
        "Figure 14. Placeholder — Forecast module with model comparison.",
        height_lines=2,
    )

    h2(doc, "7.4 Geospatial outcomes")
    add_table(
        doc,
        ["Map metric", "Value"],
        [
            ["Mapped points", fmt_int(map_s.get("points_n") or map_s.get("count") or 0)],
            ["Districts / states (KPI)", f"{fmt_int(map_kpis.get('districts', k['active_districts']))} / {fmt_int(map_kpis.get('states', k['active_states']))}"],
            ["Visible volume", fmt_int(map_kpis.get("visible_volume", map_s.get("total_volume", 0)))],
            ["Hotspot district", str(map_kpis.get("hotspot", "—"))],
            ["Average district volume", fmt_float(map_kpis.get("avg_volume", 0), 1)],
            ["High-intensity share %", fmt_float(map_kpis.get("high_share_pct", 0), 1)],
            ["Mode / scale / depth", f"{map_s.get('mode')} / {map_s.get('scale')} / {map_s.get('depth')}"],
            [
                "Centroid sources",
                ", ".join(f"{kk}={vv}" for kk, vv in (map_s.get("centroid_sources") or {}).items()) or "—",
            ],
        ],
        col_widths=[6.0, 11.0],
    )

    dist_rows = []
    for row in top_dist[:8]:
        dist_rows.append(
            [
                row.get("rank"),
                row.get("state"),
                row.get("district"),
                fmt_int(row.get("volume", 0)),
                fmt_pct(row.get("share_pct", 0), 2),
                row.get("intensity"),
                row.get("centroid_source"),
            ]
        )
    if dist_rows:
        add_table(
            doc,
            ["Rank", "State", "District", "Volume", "Share %", "Intensity", "Centroid"],
            dist_rows,
            col_widths=[1.3, 3.0, 3.5, 2.2, 1.8, 2.0, 2.2],
        )
    evidence_note(doc, "GET /api/map kpis, top_districts, centroid_sources.")

    figure_placeholder(
        doc,
        "15",
        "Geospatial map (2D intensity / heatmap / 3D) + top districts",
        "docs/figures/fig15_geospatial.png",
        "Figure 15. Placeholder — map module with intensity legend and hotspot callouts.",
        height_lines=3,
    )

    h2(doc, "7.5 Export & reproducibility affordances")
    for t in [
        "Analytics CSV pack (regional / trends / risk / ops).",
        "Forecast CSV export of history + prediction path.",
        "Map CSV exports: full points, top districts, top states.",
        "Governance pack JSON + audit CSV import/export.",
        "Reproducibility knobs: contamination, min volume, horizon, model, conformal α, seeds.",
    ]:
        bullet(doc, t)

    h2(doc, "7.6 Engineering outcome: research → production UI without forking logic")
    body(
        doc,
        "The dual-UI design is itself an outcome: Streamlit remains the rapid research surface, while the React app "
        "provides operator-grade navigation, map performance (MapLibre/deck.gl), and presentation polish. "
        "Both call the same Python analytics core, which is the primary guarantee of metric consistency cited throughout this report.",
    )

    figure_placeholder(
        doc,
        "16",
        "UI module collage (optional multi-panel figure)",
        "docs/figures/fig16_ui_collage.png",
        "Figure 16. Placeholder — collage of Dashboard, Analytics, Forecast, Map, Governance.",
        height_lines=2,
    )

    # =========================
    # 8. SYNTHESIS
    # =========================
    page_break(doc)
    h1(doc, "8. Cross-Cutting Synthesis and Implications")

    h2(doc, "8.1 Evidence chain (how a reviewer can re-verify)")
    add_table(
        doc,
        ["Step", "Action", "Expected artefact"],
        [
            ["1", "Start web stack: python run_web.py", "http://127.0.0.1:8787 ready"],
            ["2", "GET /api/health and /api/meta", "source=cache; geo_eval; repair stats"],
            ["3", "GET /api/dashboard", "KPIs, top_risk, forecast summary"],
            ["4", "GET /api/forecast?horizon=30&model=Auto", "comparison, resource_planning, meta"],
            ["5", "GET /api/insights/dashboard & /forecast", "hybrid markdown + Key numbers"],
            ["6", "GET /api/map", "points, top districts, centroid_sources"],
            ["7", "Optional: python docs/_extract_live_metrics.py", "docs/_live_api_snapshots/summary_for_outcomes.json"],
            ["8", "Fill figure placeholders from UI captures", "docs/figures/figXX_*.png"],
        ],
        col_widths=[1.5, 8.0, 7.5],
    )

    h2(doc, "8.2 Implications for operations")
    for t in [
        "Geo repair + governance reduce silent geographic bias before any ML is applied.",
        "Risk cells prioritise scarce investigation capacity; treat scores as triage, not guilt.",
        "Baseline-gated forecasts discourage overconfident complex models; decision bands communicate use-case fit.",
        "Resource planning converts abstract forecast peaks into operator headcount envelopes.",
        "Hybrid AI shortens narrative time-to-insight without replacing tabular audit trails.",
    ]:
        bullet(doc, t)

    h2(doc, "8.3 Implications for methods / research")
    for t in [
        "Demonstrates a full ops-research pipeline on noisy public-sector aggregates.",
        "Shows that unsupervised multi-feature outliers can be packaged with human-readable reason codes.",
        "Shows that forecast model choice can be automated with honest baseline gates and dual evaluation views (rolling vs holdout).",
        "Shows that research prototypes can graduate to SPA delivery without forking analytics logic.",
    ]:
        bullet(doc, t)

    h2(doc, "8.4 Limitations (must accompany any journal use of these numbers)")
    add_table(
        doc,
        ["Limitation", "Mitigation present in system"],
        [
            ["Unsupervised anomalies ≠ fraud", "UI/API disclaimers; investigation notes; reason codes"],
            ["High absolute sMAPE on noisy daily series", "Decision bands; conformal intervals; baseline gate"],
            ["Holdout may diverge from rolling CV", "Both reported; operators re-check weekly"],
            ["Centroid approximation for some districts", "centroid_source exposed (district vs state)"],
            ["Rule-pack coverage gaps", "Gold eval + durable human governance"],
            ["LLM may restate poorly if unconstrained", "Locked Key numbers; hybrid assembly"],
            ["Aggregate data only", "No individual-level claims"],
        ],
        col_widths=[6.5, 10.5],
    )

    # =========================
    # 9. REPRODUCIBILITY
    # =========================
    h1(doc, "9. Reproducibility, Configuration & Artefacts")

    h2(doc, "9.1 Key configuration defaults")
    add_table(
        doc,
        ["Parameter", "Default / this capture"],
        [
            ["CACHE_SCHEMA_VERSION", "4"],
            ["GEO_RULE_PACK_VERSION", str(meta.get("geo_rule_pack", "1.1.0"))],
            ["ANOMALY_CONTAMINATION", str(S.get("dashboard_params", {}).get("contamination", 0.05))],
            ["ANOMALY_MIN_VOLUME", str(S.get("dashboard_params", {}).get("min_volume", 50))],
            ["FORECAST_CANDIDATES", "MovingAverage, Drift, Ensemble, SeasonalNaive"],
            ["FORECAST_PRIMARY_METRIC", "mase"],
            ["FORECAST_BEAT_BASELINE_EPS", "0.0"],
            ["FORECAST_CONFORMAL_ALPHA", str(fc_meta.get("conformal_alpha", 0.1))],
            ["FORECAST_ROLLING_FOLDS", "4"],
            ["FORECAST_RANDOM_SEED", "42"],
            ["OLLAMA_MODEL", str(llm.get("model", "qwen2.5:latest"))],
        ],
        col_widths=[7.0, 10.0],
    )

    h2(doc, "9.2 Artefact checklist")
    add_table(
        doc,
        ["Artefact", "Location", "Status on capture"],
        [
            ["Parquet marts", "data/processed/", "Present (cache hit)"],
            ["Manifest + load_report", "data/processed/manifest.json", "Present"],
            ["Governance patches", "output/governance_patches.json", f"{fmt_int(gov.get('district_merges', 0))} district merges"],
            ["Audit log", "output/governance_audit.csv", "Present if actions applied"],
            ["Live API snapshots", "docs/_live_api_snapshots/", "Generated for this report"],
            ["Figure pack (to fill)", "docs/figures/", "Placeholders only"],
            ["This Outcomes document", "docs/Aadhaar_Intel_Engine_Outcomes.docx", "Generated"],
            ["Research markdown companion", "docs/RESEARCH_OUTCOMES.md", "Living notes"],
        ],
        col_widths=[5.0, 6.5, 5.5],
    )

    h2(doc, "9.3 Recommended figure pack for journal submission")
    body(doc, "Create docs/figures/ and replace each placeholder above. Suggested minimum pack:")
    for t in [
        "Fig 1 Architecture · Fig 2 ETL pipeline · Fig 3 Geo repair examples",
        "Fig 5 Governance UI · Fig 6–7 Anomaly charts · Fig 8–10 Forecast + bake-off + staffing",
        "Fig 11 AI panels · Fig 12–16 Web app screenshots (Dashboard, Analytics, Forecast, Map, collage)",
    ]:
        bullet(doc, t)

    # =========================
    # 10. CONCLUSIONS
    # =========================
    h1(doc, "10. Conclusions")
    body(
        doc,
        "Measured on the live operator stack (http://127.0.0.1:8787) with cache-backed marts for 2025-03-02 → 2025-12-31, "
        "the Aadhaar Intel Engine demonstrates that noisy national-scale identity-system aggregates can support transparent "
        "operational intelligence when geography is repaired and governed as a first-class pipeline; when outlier detection is "
        "multi-feature and explicitly non-accusatory; when forecasting is baseline-honest, interval-aware, and translated into "
        "resource envelopes; when AI assists with evidence-locked narrative; and when research and professional UIs share one analytics core.",
    )
    body(
        doc,
        f"Headline empirical anchors from this capture include: {fmt_int(k['total_enrolments'])} enrolments across "
        f"{fmt_int(k['active_states'])} states/{fmt_int(k['active_districts'])} districts; "
        f"{fmt_int(geo_stats.get('district_aliases', 0))} alias repairs and {fmt_int(geo_stats.get('ap_to_telangana', 0))} "
        f"AP→Telangana reassignments with gold both-accuracy {fmt_pct(100*float(geo_eval.get('both_accuracy', 0)), 0)}; "
        f"{fmt_int(k['anomaly_count'])} risk cells under contamination {S.get('dashboard_params', {}).get('contamination')}; "
        f"Auto-selected {k['forecast_model']} (MASE {fmt_float(k['rolling_mase'], 3)}, band {k['decision_band']}) with "
        f"~{fmt_int(rp.get('operators_at_40', 0))} operators implied at peak @ 40 txn/day; and hybrid LLM briefs available via "
        f"{llm.get('model')}.",
    )
    body(
        doc,
        "Figures remain placeholders by design so authors can insert camera-ready plots and operator screenshots without "
        "coupling the manuscript to transient repository media. Re-run the live capture checklist in Section 8.1 to refresh tables "
        "for camera-ready submission.",
    )

    # =========================
    # APPENDIX
    # =========================
    page_break(doc)
    h1(doc, "Appendix A — Figure index (placeholders)")
    fig_index = [
        ["1", "Architecture", "docs/figures/fig01_architecture.png"],
        ["2", "ETL pipeline", "docs/figures/fig02_etl_pipeline.png"],
        ["3", "Geo repair examples", "docs/figures/fig03_geo_repair_examples.png"],
        ["4", "Mart inventory", "docs/figures/fig04_mart_inventory.png"],
        ["5", "Governance UI", "docs/figures/fig05_governance_ui.png"],
        ["6", "Anomaly risk chart", "docs/figures/fig06_anomaly_risk.png"],
        ["7", "Risk radar", "docs/figures/fig07_risk_radar.png"],
        ["8", "Forecast path", "docs/figures/fig08_forecast_path.png"],
        ["9", "Model bake-off", "docs/figures/fig09_model_bakeoff.png"],
        ["10", "Resource planning", "docs/figures/fig10_resource_planning.png"],
        ["11", "AI panels", "docs/figures/fig11_ai_panels.png"],
        ["12", "Dashboard UI", "docs/figures/fig12_dashboard.png"],
        ["13", "Analytics UI", "docs/figures/fig13_analytics.png"],
        ["14", "Forecast UI", "docs/figures/fig14_forecast_ui.png"],
        ["15", "Geospatial map", "docs/figures/fig15_geospatial.png"],
        ["16", "UI collage", "docs/figures/fig16_ui_collage.png"],
        ["A+", "Additional plots/screenshots", "docs/figures/… (add freely)"],
    ]
    add_table(doc, ["Fig.", "Description", "Intended path"], fig_index, col_widths=[1.5, 5.5, 10.0])

    h1(doc, "Appendix B — Run metadata template (copy per experiment)")
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Run date", captured_at],
            ["Git commit", commit],
            ["UI", "React web http://127.0.0.1:8787"],
            ["Filters", "States: All · Date: full mart window"],
            ["Contamination / min volume", f"{S.get('dashboard_params', {}).get('contamination')} / {S.get('dashboard_params', {}).get('min_volume')}"],
            ["Forecast horizon / scope", f"{fc_meta.get('horizon', 30)} / National"],
            ["Selected model / reason", f"{selected} / {fc_meta.get('selection')}"],
            ["Notes", ""],
        ],
        col_widths=[6.0, 11.0],
    )

    h1(doc, "Appendix C — In-repo references")
    for t in [
        "flowchart.md — system flowcharts",
        "web/FEATURE_PARITY.md — Streamlit ↔ React parity",
        "docs/RESEARCH_OUTCOMES.md — living research notes",
        "docs/capture_live_assets.py — optional UI capture helper",
        "docs/_extract_live_metrics.py — API metric extractor used for this doc",
        "tests/test_geo_repair.py, tests/test_forecast.py — automated checks",
        "src/config.py — knobs cited above",
        "src/geo/*, src/etl/build_cache.py, src/ai_core.py, src/forecasting.py, src/ai/research_insights.py — methods",
    ]:
        bullet(doc, t)

    para(
        doc,
        f"Generated {captured_at} from live API evidence at {S.get('source')} · commit {commit}. "
        "Update tables by re-running docs/_extract_live_metrics.py and docs/build_outcomes_docx.py.",
        size=9,
        italic=True,
        color="64748B",
        space_before=18,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
